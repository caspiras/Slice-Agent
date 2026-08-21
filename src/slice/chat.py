"""Clean chat session for Slice IDE with tool execution capability."""

import ollama
import signal
import os
import json
import difflib
import re
import html
import ssl
import urllib.request
import urllib.error
from pathlib import Path
from rich.console import Console
from rich.live import Live
from rich.spinner import Spinner
from rich.panel import Panel
from rich.markdown import Markdown

console = Console()

# Tool-capable models (support function calling)
TOOL_CAPABLE_MODELS = [
    "llama3",
    "llama3.1",
    "llama3.2",
    "llama3.3",
    "mistral",
    "gemma",
    "gemma2",
    "gemma4",
    "command-r",
    "command-r-plus",
    "qwen",
    "qwen2",
]

# Tool definitions (standard format models expect)
TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "bash",
            "description": "Execute bash commands. Use this to: create Python/JavaScript/shell script files (use 'cat > file.py << EOF'), run Python scripts (python3 file.py), list directories (ls), search files (grep), git operations. DO NOT use this to echo answers to knowledge questions.",
            "parameters": {
                "type": "object",
                "properties": {
                    "command": {"type": "string", "description": "The bash command to execute"}
                },
                "required": ["command"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "read_document",
            "description": "Read a document file (PDF, Word, Excel, CSV, or text). Returns the complete content.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {
                        "type": "string",
                        "description": "Path to the file to read (relative to current directory)",
                    }
                },
                "required": ["file_path"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "write_document",
            "description": (
                "Write to OFFICE document files ONLY: Word (.docx), Excel (.xlsx), PowerPoint (.pptx), CSV (.csv), PDF (.pdf), text files.\n"
                "DO NOT use this for Python/JavaScript code files - use bash tool instead.\n"
                "DO NOT use this to create .app bundles - use bash to create .py files instead.\n\n"
                "SPREADSHEET EXAMPLES (Excel/CSV):\n"
                'Set cell: {"type": "set_cell", "sheet": "Sheet1", "row": 5, "col": 3, "value": "Data"}\n'
                'Add row: {"type": "append_row", "sheet": "Sheet1", "values": ["Name", "Age"]}\n'
                'Fill column: {"type": "set_column", "sheet": "Sheet1", "col": "B", "start_row": 2, "values": [10, 20]}\n\n'
                "PDF EXAMPLES:\n"
                'Add page: {"type": "add_page", "title": "Page Title", "content": "Page content"}\n'
                'Add paragraph: {"type": "add_paragraph", "text": "Paragraph text", "font_size": 12}\n'
                'Add text: {"type": "add_text", "text": "Text content", "font_size": 14}\n\n'
                "Multiple operations - use array:\n"
                '[{"type": "add_page", "title": "Intro"}, {"type": "add_paragraph", "text": "Content"}]'
            ),
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path to the file to write"},
                    "operations": {
                        "type": "string",
                        "description": "JSON string of operation(s). Single object or array of objects. Must be valid JSON.",
                    },
                },
                "required": ["file_path", "operations"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "edit_code",
            "description": "Edit a source code file by replacing old content with new content. Shows a diff for user approval before applying. Use this for code files (.py, .js, .java, etc.). For Office documents use write_document instead.",
            "parameters": {
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path to the code file to edit"},
                    "old_content": {
                        "type": "string",
                        "description": "The exact text to find and replace (must match exactly including whitespace)",
                    },
                    "new_content": {
                        "type": "string",
                        "description": "The new text to replace it with",
                    },
                    "description": {
                        "type": "string",
                        "description": "Brief description of what this edit does (e.g., 'Fix typo in function name', 'Add error handling')",
                    },
                },
                "required": ["file_path", "old_content", "new_content", "description"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "convert_to_json",
            "description": "Convert document files to JSON format efficiently (handles large files). Supports Excel (.xlsx), CSV (.csv), Word (.docx with tables), and PDF (.pdf). Uses chunking/streaming to avoid memory issues.",
            "parameters": {
                "type": "object",
                "properties": {
                    "input_file": {
                        "type": "string",
                        "description": "Path to the input file to convert",
                    },
                    "output_file": {
                        "type": "string",
                        "description": "Path where the JSON output should be saved",
                    },
                },
                "required": ["input_file", "output_file"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "convert_to_markdown",
            "description": "Convert document files to Markdown format efficiently (handles large files). Supports Excel (.xlsx), CSV (.csv), Word (.docx with tables), and PDF (.pdf). Tables are converted to Markdown table syntax.",
            "parameters": {
                "type": "object",
                "properties": {
                    "input_file": {
                        "type": "string",
                        "description": "Path to the input file to convert",
                    },
                    "output_file": {
                        "type": "string",
                        "description": "Path where the Markdown output should be saved (.md extension)",
                    },
                },
                "required": ["input_file", "output_file"],
            },
        },
    },
    {
        "type": "function",
        "function": {
            "name": "fetch_url",
            "description": "Fetch and read the text content of a web page or URL (http:// or https://). Use this whenever the user asks you to review, read, summarize, check, or look at a URL, link, or web page. Returns the page text. Requires user permission before fetching. DO NOT say you cannot access URLs - call this tool instead.",
            "parameters": {
                "type": "object",
                "properties": {
                    "url": {
                        "type": "string",
                        "description": "The full http:// or https:// URL to fetch",
                    }
                },
                "required": ["url"],
            },
        },
    },
]


class ChatSession:
    """Chat session with an Ollama model - clean IDE experience with tool execution."""

    # Constants
    MAX_DOCUMENT_CHARS = 100000  # ~100KB max for document content
    CSV_CHUNK_SIZE = 10000  # Process CSV files in 10k row chunks
    URL_FETCH_TIMEOUT = 30  # Timeout in seconds for fetching a URL
    MAX_TOOL_ROUNDS = 10  # Max sequential tool-call rounds per user turn (loop guard)
    PROJECT_INSTRUCTIONS_FILE = "SLICE.md"  # Auto-loaded per-project instructions

    def __init__(self, model_name: str, safe_directory: str, skill_loader=None):
        self.model_name = model_name
        self.safe_directory = safe_directory
        self.interrupted = False
        self.skill_loader = skill_loader

        # Always try tools - let Ollama decide if model supports them
        self.supports_tools = True

        # Initialize conversation with system guidance
        self.conversation_history = [
            {
                "role": "system",
                "content": (
                    "You are a helpful assistant with capabilities to read/write documents, edit code, and execute bash commands.\n\n"
                    "CRITICAL - Creating Python/JavaScript Apps:\n"
                    "When user asks to create a Python app (.py file):\n"
                    "1. Use bash tool with: cat > filename.py << EOF\\n<code>\\nEOF\n"
                    "2. Then use bash tool again with: python3 filename.py\n"
                    "NEVER use write_document for .py files - it's only for Office documents (Word, Excel, PDF).\n"
                    "NEVER try to create .app bundles - create .py files instead.\n"
                    "ALWAYS create the file BEFORE trying to run it.\n\n"
                    "CRITICAL - You have TOOLS that you must CALL, do NOT just talk about using them:\n"
                    "- bash tool - Creates/runs Python/JS files, executes commands\n"
                    "- read_document tool - Reads Office documents and text files\n"
                    "- write_document tool - Writes Office documents ONLY (Word, Excel, PDF, PowerPoint, CSV)\n"
                    "- edit_code tool - Edits existing source code files\n"
                    "- fetch_url tool - Fetches/reads the contents of a web page or URL\n"
                    "IMPORTANT: You can call MULTIPLE tools in a SINGLE response. For example, call bash twice to create and then run a file.\n\n"
                    "CRITICAL - Action vs. Explanation:\n"
                    "- When user asks you to CREATE, MAKE, BUILD, WRITE, RUN, or EXECUTE something → CALL THE TOOL NOW\n"
                    "- 'Yes' or 'Y' or 'Sure' or 'OK' or 'Go ahead' → If context is running a file, CALL bash tool with 'python3 <filename>'\n"
                    "- 'run the app' or 'execute X' or 'please run X' → CALL bash tool NOW with appropriate command\n"
                    "- When user asks HOW to do something or WHAT is something → Only then explain with text\n\n"
                    "CRITICAL - After Creating Executable Files:\n"
                    "When user asks you to create an app/script (.py, .js, .sh files), respond with bash tool call to create AND run in sequence.\n"
                    "After the creation tool call result comes back, your NEXT response should be ANOTHER bash tool call to run it.\n"
                    "DO NOT respond with text after creating the file - immediately call bash again to run it.\n\n"
                    "EXAMPLE CONVERSATION FLOW:\n"
                    "User: 'create a Python app'\n"
                    "Assistant: [calls bash tool: cat > app.py << EOF...]\n"
                    "System: [returns result: Command executed successfully]\n"
                    "Assistant: [calls bash tool again: python3 app.py] ← THIS IS YOUR RESPONSE, not text\n"
                    "System: [returns result with output or error]\n"
                    "Assistant: [NOW you can respond with text: Created and ran app.py]\n\n"
                    "User: 'run the app' → You CALL bash tool with command 'python3 app.py'\n"
                    "User: 'Yes' (when previous context was about running) → You CALL bash tool with 'python3 <filename>'\n\n"
                    "WRONG BEHAVIOR - NEVER DO THIS:\n"
                    "User: 'create an app' → You ONLY call bash('python3 app.py') without creating the file first ← WRONG! Create file FIRST, then run.\n"
                    "User: 'create an app' → You call bash to create file, then WAIT. ← WRONG! Call bash AGAIN in the same response to run it.\n"
                    "User: 'create an app' → You create file, then respond 'Created app.py. Would you like me to run it?' ← WRONG! Make the second bash call.\n"
                    "User: 'Yes' → You respond: 'Running app.py' ← WRONG! You must CALL the bash tool!\n\n"
                    "After executing a tool successfully, respond concisely - DO NOT explain the code or give instructions\n"
                    "The user already sees each tool's raw output in the terminal. Do NOT paste that raw output back\n"
                    "verbatim. Instead add value: present it usefully, summarize, or answer their question. If a project's\n"
                    "SLICE.md specifies an exact display/formatting for a command's data, follow that. If the output already\n"
                    "answers the request and there is nothing to add, reply with a brief one-line confirmation.\n"
                    "Default to ACTION (calling tools) when request is ambiguous - user can always deny the permission prompt\n\n"
                    "IMPORTANT - When to use tools:\n"
                    "- Use bash tool for file/system operations (create files, list directories, git commands, etc.)\n"
                    "- Use read_document tool to read PDF, Word, Excel, CSV, and text/code files - DO NOT verify file existence with ls first, just read it directly\n"
                    "- Use edit_code tool to modify source code files (.py, .js, .java, etc.) - shows a diff for user approval\n"
                    "- Use write_document tool for ALL document types (Word, Excel, PowerPoint, CSV, PDF, text)\n"
                    "- For general knowledge questions, answer directly with text - DO NOT use bash to echo answers\n\n"
                    "Code editing workflow:\n"
                    "1. Read the file with read_document\n"
                    "2. Identify the section to change\n"
                    "3. Use edit_code with exact old_content and new_content\n"
                    "4. The diff will be shown to user for approval\n\n"
                    "Spreadsheet editing workflow (Excel .xlsx, CSV .csv):\n"
                    "1. ALWAYS read the file first with read_document to see current structure\n"
                    "2. Identify what needs to change (which rows, columns, cells)\n"
                    "3. Use write_document with JSON operations\n"
                    "4. Common operations:\n"
                    '   - Set specific cell: {"type": "set_cell", "sheet": "Sheet1", "row": 2, "col": "A", "value": "Data"}\n'
                    '   - Add new row: {"type": "append_row", "sheet": "Sheet1", "values": ["col1", "col2", "col3"]}\n'
                    '   - Fill column: {"type": "set_column", "sheet": "Sheet1", "col": "B", "start_row": 2, "values": [10, 20, 30]}\n'
                    '5. For CSV files, omit the "sheet" parameter\n'
                    "6. Multiple operations can be combined in an array: [{...}, {...}]\n"
                    '7. Columns can be letters ("A", "M") or numbers (1, 13)\n'
                    "8. Rows are 1-indexed (row 1 is first row)\n\n"
                    "PDF editing workflow (.pdf):\n"
                    "1. PDFs can be created and edited with write_document\n"
                    "2. Common operations:\n"
                    '   - Add page with title and content: {"type": "add_page", "title": "Page Title", "content": "Content"}\n'
                    '   - Add paragraph: {"type": "add_paragraph", "text": "Text content", "font_size": 12}\n'
                    '   - Add text: {"type": "add_text", "text": "Text", "font_size": 14}\n'
                    "3. Multiple pages/paragraphs can be combined in an array\n"
                    "4. PDFs are built sequentially - operations are applied in order\n\n"
                    "Git operations (use bash tool - all require user approval via permission prompt):\n"
                    "- Read-only operations (safe to suggest): git status, git log, git diff, git show, git branch\n"
                    "- Local operations (safe to suggest): git add, git commit, git checkout -b, git merge\n"
                    "- Remote operations: NEVER suggest git push or git pull unless user EXPLICITLY asks in their message\n"
                    "- After making local commits, remind user they can push when ready, don't run push automatically\n\n"
                    "File format conversion - Use convert_to_json or convert_to_markdown tools:\n"
                    "- Use convert_to_json for Excel (.xlsx), CSV (.csv), Word (.docx), and PDF (.pdf) files\n"
                    "- Use convert_to_markdown to convert documents to Markdown (.md) format with table support\n"
                    "- Both tools handle large files efficiently with chunking/streaming to avoid memory errors\n"
                    "- Word documents: Extracts BOTH paragraphs AND tables (tables were missing before!)\n"
                    "- CSV/Excel to Markdown: Converts tables to Markdown table syntax with | separators\n"
                    "- CSV files: Uses chunking for large files (processes 10k rows at a time)\n"
                    "- Excel files: Converts all rows to JSON array or Markdown tables\n"
                    "- PDF files: Extracts text page-by-page to avoid memory issues\n"
                    "- Example JSON: convert_to_json with input_file='data.xlsx' and output_file='data.json'\n"
                    "- Example Markdown: convert_to_markdown with input_file='data.xlsx' and output_file='data.md'\n"
                    "- For JSON files, treat them as text files with write_document using replace_content operation\n\n"
                    "File operations:\n"
                    "- When user asks about file content (not conversion), read it directly with read_document\n"
                    "- Don't use ls or find to verify files exist before reading them\n"
                    "- Use bash 'touch filename' to create empty text files only\n"
                    "- NEVER try to create empty document files (Excel, Word, PowerPoint, PDF) with touch - they need proper structure\n"
                    "- To create new document files, use write_document with operations (it will create the file with proper structure)\n"
                    "- Multi-file search: use bash with grep -r, find, or other search tools\n\n"
                    "Web / URL access (use fetch_url tool - requires user approval via permission prompt):\n"
                    "- When the user gives you a URL or asks you to review, read, summarize, or check a web page, CALL fetch_url with that URL\n"
                    "- NEVER say you cannot access the internet or URLs - you CAN, via the fetch_url tool (the user approves each fetch)\n"
                    "- Only http:// and https:// URLs are supported\n"
                    "- After fetching, summarize or answer based on the returned page text\n\n"
                    "Language:\n"
                    "- Always respond in English unless the user writes to you in another language\n\n"
                    "Formatting guidelines:\n"
                    "- Use simple paragraph formatting with bullet points (•) or numbered lists\n"
                    "- Avoid markdown tables with | symbols - they display poorly in terminals\n"
                    "- Use simple text formatting instead of complex markdown"
                ),
            }
        ]

        # Load per-project instructions from SLICE.md (if present) as a second
        # system message. This gives the model persistent, project-specific
        # guidance as soon as the session starts (like CLAUDE.md for Claude Code).
        # The instructions are auto-reloaded when the file changes (see
        # refresh_project_instructions), so editing SLICE.md mid-session takes
        # effect on the next prompt without a restart.
        self.has_project_instructions = False
        self._project_instructions_text = ""  # last-synced SLICE.md content
        self.refresh_project_instructions()

    def _project_instructions_header(self) -> str:
        """The fixed prefix prepended to SLICE.md content in the system message.

        Also used to locate the instruction message in the conversation history,
        so building and detection can never drift apart.
        """
        return (
            "The user has provided project-specific instructions in a file named "
            f"{self.PROJECT_INSTRUCTIONS_FILE}. Treat these as authoritative guidance "
            "for how to work in this project, and follow them carefully:\n\n"
        )

    def _read_project_instructions(self) -> str:
        """Read SLICE.md from the sandbox directory, if it exists.

        Returns the file's text content, or an empty string when the file is
        absent or unreadable. Never raises - a missing/broken SLICE.md must not
        prevent a session from starting.
        """
        path = os.path.join(self.safe_directory, self.PROJECT_INSTRUCTIONS_FILE)
        try:
            if not os.path.isfile(path):
                return ""
            with open(path, "r", encoding="utf-8") as f:
                return f.read().strip()
        except Exception:
            return ""

    def refresh_project_instructions(self, force: bool = False) -> None:
        """Reload SLICE.md and sync the injected system message in place.

        Handles create / edit / delete of SLICE.md between calls. Safe to call
        every turn - it does nothing when the content is unchanged (unless
        force=True, used after a model switch where the copied history may carry
        a stale copy of the instructions).
        """
        text = self._read_project_instructions()

        if not force and text == self._project_instructions_text:
            return

        # Locate an existing SLICE.md system message by its fixed header prefix.
        header = self._project_instructions_header()
        idx = None
        for i, msg in enumerate(self.conversation_history):
            if (
                msg.get("role") == "system"
                and isinstance(msg.get("content"), str)
                and msg["content"].startswith(header)
            ):
                idx = i
                break

        if text:
            new_msg = {"role": "system", "content": header + text}
            if idx is not None:
                self.conversation_history[idx] = new_msg
            else:
                # Insert right after the base system message (index 0), if any.
                insert_at = (
                    1
                    if self.conversation_history
                    and self.conversation_history[0].get("role") == "system"
                    else 0
                )
                self.conversation_history.insert(insert_at, new_msg)
            self.has_project_instructions = True
        else:
            # SLICE.md was removed or emptied - drop the instruction message.
            if idx is not None:
                del self.conversation_history[idx]
            self.has_project_instructions = False

        self._project_instructions_text = text

    def _execute_command(self, command: str) -> str:
        """Execute a command with user permission."""
        from .executor import CommandExecutor

        # Use Python executor for sandboxing and permission prompts
        executor = CommandExecutor(self.safe_directory)
        result = executor.execute_with_permission(
            command, context="Model requested command execution"
        )

        if result.get("cancelled"):
            return "Command cancelled by user."
        elif result.get("success"):
            output = result.get("output", "").strip()
            return f"Command executed successfully.{f' Output: {output}' if output else ''}"
        else:
            error = result.get("error", "Unknown error")
            return f"Command failed: {error}"

    def _read_document(self, file_path: str) -> str:
        """Read a document file."""
        from .document_reader import read_document

        # Resolve path relative to safe directory
        full_path = os.path.join(self.safe_directory, file_path)

        try:
            result = read_document(full_path)

            if result.get("success"):
                # Return the content from the dict
                content = result.get("content", "")
                file_type = result.get("file_type", "unknown")

                # Warn and truncate if content is very large
                if len(content) > self.MAX_DOCUMENT_CHARS:
                    console.print(
                        f"[yellow]⚠️  Large document ({len(content)} chars) - truncating to first {self.MAX_DOCUMENT_CHARS} chars[/yellow]"
                    )
                    content = (
                        content[: self.MAX_DOCUMENT_CHARS]
                        + f"\n\n[... truncated {len(content) - self.MAX_DOCUMENT_CHARS} additional characters ...]"
                    )

                return f"[{file_type} file content]\n{content}"
            else:
                # Return the error message
                error = result.get("error", "Unknown error")
                return f"Error: {error}"

        except Exception as e:
            return f"Error reading document: {str(e)}"

    def _write_document(self, file_path: str, operations: str) -> str:
        """Write to a document file."""
        from .document_writer import write_document

        # Resolve path relative to safe directory
        full_path = os.path.join(self.safe_directory, file_path)

        try:
            # Parse operations JSON
            ops = json.loads(operations)
            result = write_document(full_path, ops)

            if result.get("success"):
                # Return success message
                message = result.get("message", "Document updated successfully")
                ops_count = result.get("operations_applied", 0)
                return f"{message} ({ops_count} operations applied)"
            else:
                # Return error message
                error = result.get("error", "Unknown error")
                return f"Error: {error}"

        except json.JSONDecodeError:
            return "Error: Invalid JSON in operations parameter"
        except Exception as e:
            return f"Error writing document: {str(e)}"

    def _edit_code(
        self, file_path: str, old_content: str, new_content: str, description: str
    ) -> str:
        """Edit a code file with diff preview and user approval."""
        # Resolve path relative to safe directory
        full_path = os.path.join(self.safe_directory, file_path)

        try:
            # Read current file content
            if not os.path.exists(full_path):
                return f"Error: File not found: {file_path}"

            with open(full_path, "r", encoding="utf-8") as f:
                current_content = f.read()

            # Check if old_content exists in file
            if old_content not in current_content:
                return f"Error: Could not find the specified content to replace in {file_path}. Make sure old_content matches exactly."

            # Generate new file content
            updated_content = current_content.replace(old_content, new_content, 1)

            # Show diff
            diff = difflib.unified_diff(
                current_content.splitlines(keepends=True),
                updated_content.splitlines(keepends=True),
                fromfile=f"{file_path} (current)",
                tofile=f"{file_path} (proposed)",
                lineterm="",
            )
            diff_text = "".join(diff)

            # Display the edit request
            console.print("\n[bold cyan]📝 Code Edit Request[/bold cyan]")
            console.print(f"[dim]{description}[/dim]\n")

            # Show diff in a panel
            from rich.syntax import Syntax

            diff_syntax = Syntax(diff_text, "diff", theme="monokai", line_numbers=False)
            console.print(Panel(diff_syntax, title=f"Changes to {file_path}", border_style="cyan"))

            # Ask for permission
            response = input("\nApply these changes? (y/N): ").strip().lower()

            if response == "y":
                # Write the updated content
                with open(full_path, "w", encoding="utf-8") as f:
                    f.write(updated_content)
                console.print(f"[green]✓ Changes applied to {file_path}[/green]\n")
                return f"Successfully edited {file_path}: {description}"
            else:
                console.print("[yellow]✗ Changes cancelled[/yellow]\n")
                return "Edit cancelled by user"

        except Exception as e:
            return f"Error editing file: {str(e)}"

    def _convert_to_json(self, input_file: str, output_file: str) -> str:
        """Convert document files to JSON format with chunking for large files."""
        # Resolve paths relative to safe directory
        input_path = os.path.join(self.safe_directory, input_file)
        output_path = os.path.join(self.safe_directory, output_file)

        if not os.path.exists(input_path):
            return f"Error: Input file not found: {input_file}"

        suffix = Path(input_path).suffix.lower()

        try:
            if suffix == ".xlsx":
                return self._convert_excel_to_json(input_path, output_path)
            elif suffix == ".csv":
                return self._convert_csv_to_json(input_path, output_path)
            elif suffix == ".docx":
                return self._convert_word_to_json(input_path, output_path)
            elif suffix == ".pdf":
                return self._convert_pdf_to_json(input_path, output_path)
            else:
                return (
                    f"Error: Unsupported file type: {suffix}. Supported: .xlsx, .csv, .docx, .pdf"
                )
        except Exception as e:
            return f"Error converting {input_file} to JSON: {str(e)}"

    def _convert_excel_to_json(self, input_path: str, output_path: str) -> str:
        """Convert Excel file to JSON."""
        import pandas as pd

        df = pd.read_excel(input_path, engine="openpyxl")
        result = df.to_dict(orient="records")

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=2, ensure_ascii=False, default=str)

        return f"Successfully converted {len(result)} rows from Excel to JSON: {output_path}"

    def _convert_csv_to_json(self, input_path: str, output_path: str) -> str:
        """Convert CSV file to JSON with chunking for large files."""
        import pandas as pd

        # Read in chunks to handle large files
        chunk_size = self.CSV_CHUNK_SIZE
        chunks = []

        for chunk in pd.read_csv(input_path, chunksize=chunk_size, encoding="utf-8"):
            chunks.append(chunk)

        # Combine all chunks
        df = pd.concat(chunks, ignore_index=True)
        result = df.to_dict(orient="records")

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=2, ensure_ascii=False, default=str)

        return f"Successfully converted {len(result)} rows from CSV to JSON: {output_path}"

    def _convert_word_to_json(self, input_path: str, output_path: str) -> str:
        """Convert Word document to JSON, including tables."""
        from docx import Document

        doc = Document(input_path)
        result = {"paragraphs": [], "tables": []}

        # Extract paragraphs and tables in order
        for element in doc.element.body:
            if element.tag.endswith("p"):
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if text:
                            result["paragraphs"].append(text)
                        break
            elif element.tag.endswith("tbl"):
                for table in doc.tables:
                    if table._element == element:
                        table_data = []
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            if any(row_data):
                                table_data.append(row_data)
                        if table_data:
                            result["tables"].append(table_data)
                        break

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

        return f"Successfully converted Word document to JSON: {len(result['paragraphs'])} paragraphs, {len(result['tables'])} tables → {output_path}"

    def _convert_pdf_to_json(self, input_path: str, output_path: str) -> str:
        """Convert PDF to JSON, processing page by page."""
        from pypdf import PdfReader

        reader = PdfReader(input_path)
        result = {"pages": [], "metadata": {"page_count": len(reader.pages)}}

        # Process pages one at a time to avoid memory issues
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()
            result["pages"].append({"page_number": page_num, "text": page_text})

        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(result, f, indent=2, ensure_ascii=False)

        return (
            f"Successfully converted {len(result['pages'])} pages from PDF to JSON: {output_path}"
        )

    def _convert_to_markdown(self, input_file: str, output_file: str) -> str:
        """Convert document files to Markdown format."""

        # Resolve paths relative to safe directory
        input_path = os.path.join(self.safe_directory, input_file)
        output_path = os.path.join(self.safe_directory, output_file)

        if not os.path.exists(input_path):
            return f"Error: Input file not found: {input_file}"

        suffix = Path(input_path).suffix.lower()

        try:
            if suffix == ".xlsx":
                return self._convert_excel_to_markdown(input_path, output_path)
            elif suffix == ".csv":
                return self._convert_csv_to_markdown(input_path, output_path)
            elif suffix == ".docx":
                return self._convert_word_to_markdown(input_path, output_path)
            elif suffix == ".pdf":
                return self._convert_pdf_to_markdown(input_path, output_path)
            else:
                return (
                    f"Error: Unsupported file type: {suffix}. Supported: .xlsx, .csv, .docx, .pdf"
                )
        except Exception as e:
            return f"Error converting {input_file} to Markdown: {str(e)}"

    def _convert_excel_to_markdown(self, input_path: str, output_path: str) -> str:
        """Convert Excel file to Markdown with table formatting."""
        import pandas as pd

        # Read Excel file
        df = pd.read_excel(input_path, engine="openpyxl")

        # Convert to Markdown table
        md_content = df.to_markdown(index=False)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md_content)

        return f"Successfully converted {len(df)} rows from Excel to Markdown table: {output_path}"

    def _convert_csv_to_markdown(self, input_path: str, output_path: str) -> str:
        """Convert CSV file to Markdown with table formatting and chunking for large files."""
        import pandas as pd

        # Read in chunks to handle large files
        chunk_size = self.CSV_CHUNK_SIZE
        chunks = []

        for chunk in pd.read_csv(input_path, chunksize=chunk_size, encoding="utf-8"):
            chunks.append(chunk)

        # Combine all chunks
        df = pd.concat(chunks, ignore_index=True)

        # Convert to Markdown table
        md_content = df.to_markdown(index=False)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(md_content)

        return f"Successfully converted {len(df)} rows from CSV to Markdown table: {output_path}"

    def _convert_word_to_markdown(self, input_path: str, output_path: str) -> str:
        """Convert Word document to Markdown, including tables."""
        from docx import Document

        doc = Document(input_path)
        md_lines = []

        # Extract paragraphs and tables in order
        for element in doc.element.body:
            if element.tag.endswith("p"):
                for para in doc.paragraphs:
                    if para._element == element:
                        text = para.text.strip()
                        if text:
                            md_lines.append(text)
                            md_lines.append("")  # Blank line after paragraph
                        break
            elif element.tag.endswith("tbl"):
                for table in doc.tables:
                    if table._element == element:
                        # Convert table to Markdown format
                        table_rows = []
                        for i, row in enumerate(table.rows):
                            row_data = [cell.text.strip() for cell in row.cells]
                            table_rows.append("| " + " | ".join(row_data) + " |")

                            # Add separator after header row
                            if i == 0:
                                separator = "| " + " | ".join(["---"] * len(row_data)) + " |"
                                table_rows.append(separator)

                        md_lines.extend(table_rows)
                        md_lines.append("")  # Blank line after table
                        break

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_lines))

        return f"Successfully converted Word document to Markdown: {output_path}"

    def _convert_pdf_to_markdown(self, input_path: str, output_path: str) -> str:
        """Convert PDF to Markdown, processing page by page."""
        from pypdf import PdfReader

        reader = PdfReader(input_path)
        md_lines = []

        # Process pages one at a time to avoid memory issues
        for page_num, page in enumerate(reader.pages, start=1):
            page_text = page.extract_text()
            md_lines.append(f"## Page {page_num}\n")
            md_lines.append(page_text)
            md_lines.append("\n---\n")  # Page separator

        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(md_lines))

        return (
            f"Successfully converted {len(reader.pages)} pages from PDF to Markdown: {output_path}"
        )

    def _html_to_text(self, raw_html: str) -> str:
        """Convert HTML to readable plain text (dependency-free).

        Strips <script>/<style> blocks, removes remaining tags, unescapes
        HTML entities, and collapses excess whitespace. Good enough for the
        model to review/summarize a page - not a full HTML renderer.
        """
        # Remove script and style blocks entirely (including their content)
        text = re.sub(
            r"<(script|style)\b[^>]*>.*?</\1>", " ", raw_html, flags=re.DOTALL | re.IGNORECASE
        )
        # Treat block-level break tags as line breaks for readability
        text = re.sub(r"(?i)<(br|/p|/div|/li|/h[1-6]|/tr)\s*/?>", "\n", text)
        # Strip all remaining tags
        text = re.sub(r"<[^>]+>", " ", text)
        # Unescape entities (&amp; &nbsp; etc.)
        text = html.unescape(text)
        # Collapse runs of spaces/tabs, and trim blank lines
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)
        return text.strip()

    def _fetch_url(self, url: str) -> str:
        """Fetch a URL after user permission and return its text content."""
        url = url.strip()
        if not (url.startswith("http://") or url.startswith("https://")):
            return "Error: Only http:// and https:// URLs are supported."

        # Permission prompt (mirrors the edit_code approval style)
        console.print("\n[bold yellow]🌐 Web Request[/bold yellow]")
        console.print("[dim]Model requested to fetch a URL[/dim]")
        console.print(Panel(url, title="URL", border_style="yellow"))

        try:
            response = input("Fetch this URL? (y/N): ").strip().lower()
        except (KeyboardInterrupt, EOFError):
            console.print("\n[dim]Action cancelled by user[/dim]\n")
            return "URL fetch cancelled by user."

        if response not in ("y", "yes"):
            console.print("[dim]Action cancelled by user[/dim]\n")
            return "URL fetch cancelled by user."

        # Build an SSL context with a real CA bundle. Some Python installs
        # (notably python.org builds on macOS) ship without wired-up system
        # certificates, which makes HTTPS fail with CERTIFICATE_VERIFY_FAILED.
        # Prefer certifi's bundle when available, otherwise fall back to the
        # default context.
        try:
            import certifi

            ssl_context = ssl.create_default_context(cafile=certifi.where())
        except Exception:
            ssl_context = ssl.create_default_context()

        # Fetch with a spinner
        try:
            request = urllib.request.Request(
                url, headers={"User-Agent": "Slice/1.6 (local Ollama IDE)"}
            )
            with Live(
                Spinner("dots", text=f"[cyan]fetching {url}...[/cyan]"),
                console=console,
                transient=True,
            ) as fetch_live:
                with urllib.request.urlopen(
                    request, timeout=self.URL_FETCH_TIMEOUT, context=ssl_context
                ) as resp:
                    content_type = resp.headers.get_content_type()
                    charset = resp.headers.get_content_charset() or "utf-8"
                    raw = resp.read()
                fetch_live.stop()
        except urllib.error.HTTPError as e:
            console.print(f"[red]✗ HTTP error {e.code} fetching URL[/red]\n")
            return f"Error fetching URL: HTTP {e.code} {e.reason}"
        except Exception as e:
            console.print("[red]✗ Failed to fetch URL[/red]\n")
            return f"Error fetching URL: {str(e)}"

        text = raw.decode(charset, errors="replace")
        if "html" in content_type:
            text = self._html_to_text(text)

        if len(text) > self.MAX_DOCUMENT_CHARS:
            console.print(
                f"[yellow]⚠️  Large page ({len(text)} chars) - truncating to first {self.MAX_DOCUMENT_CHARS} chars[/yellow]"
            )
            text = (
                text[: self.MAX_DOCUMENT_CHARS]
                + f"\n\n[... truncated {len(text) - self.MAX_DOCUMENT_CHARS} additional characters ...]"
            )

        console.print(f"[green]✓ Fetched {url}[/green]\n")
        return f"[content from {url} ({content_type})]\n{text}"

    def _stream_assistant_text(self, stream, spinner_live=None):
        """Consume an Ollama chat stream and render the assistant text as Markdown.

        Feeding the accumulated text through rich.markdown.Markdown means headers,
        bold, lists and tables render formatted instead of showing as raw
        '**' / '#' / '|' syntax. A Live region re-renders on each token so the
        response still appears progressively, and the formatted output persists
        once streaming completes.

        `spinner_live`, if given, is the active "baking..." spinner; it is stopped
        as soon as the first content or tool call arrives. Returns
        (response_text, tool_calls). Respects self.interrupted for early exit and
        strips <tool_call> tags that some models emit.
        """
        response_text = ""
        tool_calls = None

        def stop_spinner():
            if spinner_live is not None and spinner_live.is_started:
                spinner_live.stop()

        # Non-transient Live so the rendered Markdown stays on screen afterward.
        md_live = Live(console=console, refresh_per_second=15, transient=False)
        md_started = False
        try:
            for chunk in stream:
                if self.interrupted:
                    break

                message = chunk.get("message", {})

                # Tool calls can arrive in any chunk, before done=True
                chunk_tool_calls = message.get("tool_calls")
                if chunk_tool_calls:
                    tool_calls = chunk_tool_calls
                    stop_spinner()

                if chunk.get("done"):
                    break

                content = message.get("content", "")
                if content and "<tool_call>" not in content and "</tool_call>" not in content:
                    stop_spinner()
                    response_text += content
                    if not md_started:
                        md_live.start()
                        md_started = True
                    md_live.update(Markdown(response_text))
        finally:
            if md_started:
                md_live.stop()
            stop_spinner()

        if self.interrupted:
            console.print("\n[yellow]⚠️  Generation interrupted[/yellow]")

        return response_text, tool_calls

    def _dispatch_tool_call(self, name: str, arguments: dict) -> str:
        """Execute a single tool call and return its result string for the model.

        Each tool renders its own UI (permission prompts, spinners) as needed.
        Empty/invalid parameters return an error string (instead of silently
        skipping) so the model gets feedback and can correct itself next round.
        """
        if name == "bash":
            command = arguments.get("command", "")
            if not command:
                console.print("[red]Error: Model provided empty command[/red]")
                return "Error: empty command provided"
            # Execute the command (has its own permission UI)
            return self._execute_command(command)

        if name == "read_document":
            file_path = arguments.get("file_path", "")
            if not file_path:
                console.print("[red]Error: Model provided empty file path[/red]")
                return "Error: empty file path provided"
            # Show spinner while reading (can take time for large files)
            with Live(
                Spinner("dots", text=f"[cyan]reading {file_path}...[/cyan]"),
                console=console,
                transient=True,
            ) as read_live:
                result = self._read_document(file_path)
                read_live.stop()
            console.print("[green]✓ Document loaded[/green]\n")
            return result

        if name == "write_document":
            file_path = arguments.get("file_path", "")
            operations = arguments.get("operations", "")
            if not file_path or not operations:
                console.print("[red]Error: Model provided incomplete parameters[/red]")
                return "Error: incomplete parameters (need file_path and operations)"
            with Live(
                Spinner("dots", text=f"[cyan]writing {file_path}...[/cyan]"),
                console=console,
                transient=True,
            ) as write_live:
                result = self._write_document(file_path, operations)
                write_live.stop()
            if result.startswith("Error:"):
                console.print(f"[red]✗ {result}[/red]\n")
            else:
                console.print("[green]✓ Document updated[/green]\n")
            return result

        if name == "edit_code":
            file_path = arguments.get("file_path", "")
            old_content = arguments.get("old_content", "")
            new_content = arguments.get("new_content", "")
            description = arguments.get("description", "")
            if not file_path or not old_content or not new_content:
                console.print("[red]Error: Model provided incomplete parameters[/red]")
                return "Error: incomplete parameters (need file_path, old_content, new_content)"
            # Edit the code file (shows diff and asks for permission)
            return self._edit_code(file_path, old_content, new_content, description)

        if name == "convert_to_json":
            input_file = arguments.get("input_file", "")
            output_file = arguments.get("output_file", "")
            if not input_file or not output_file:
                console.print("[red]Error: Model provided incomplete parameters[/red]")
                return "Error: incomplete parameters (need input_file and output_file)"
            with Live(
                Spinner("dots", text=f"[cyan]converting {input_file} to JSON...[/cyan]"),
                console=console,
                transient=True,
            ) as convert_live:
                result = self._convert_to_json(input_file, output_file)
                convert_live.stop()
            if result.startswith("Error:"):
                console.print(f"[red]✗ {result}[/red]\n")
            else:
                console.print(f"[green]✓ {result}[/green]\n")
            return result

        if name == "convert_to_markdown":
            input_file = arguments.get("input_file", "")
            output_file = arguments.get("output_file", "")
            if not input_file or not output_file:
                console.print("[red]Error: Model provided incomplete parameters[/red]")
                return "Error: incomplete parameters (need input_file and output_file)"
            with Live(
                Spinner("dots", text=f"[cyan]converting {input_file} to Markdown...[/cyan]"),
                console=console,
                transient=True,
            ) as convert_live:
                result = self._convert_to_markdown(input_file, output_file)
                convert_live.stop()
            if result.startswith("Error:"):
                console.print(f"[red]✗ {result}[/red]\n")
            else:
                console.print(f"[green]✓ {result}[/green]\n")
            return result

        if name == "fetch_url":
            url = arguments.get("url", "")
            if not url:
                console.print("[red]Error: Model provided empty URL[/red]")
                return "Error: empty URL provided"
            # Fetch the URL (asks permission, shows its own UI)
            return self._fetch_url(url)

        console.print(f"[yellow]⚠️  Model requested unknown tool: {name}[/yellow]")
        return f"Error: unknown tool '{name}'"

    def process_stream(self, user_input: str):
        """
        Process user input and stream response from model.
        Supports tool calling for models that can use it.
        Returns True if completed, False if interrupted.
        """
        # Pick up any edits to SLICE.md before responding (auto-reload).
        self.refresh_project_instructions()

        # Check if this is a skill command
        if user_input.strip().startswith("/") and self.skill_loader:
            # Extract just the skill name (first word after /)
            skill_name = user_input.strip()[1:].split()[0] if user_input.strip()[1:] else ""
            skill = self.skill_loader.get_skill(skill_name)

            if skill:
                console.print(f"[cyan]🔧 Running skill: {skill.name}[/cyan]")
                console.print(f"[dim]{skill.description}[/dim]\n")

                # Inject skill instructions as a system message
                self.conversation_history.append(
                    {
                        "role": "system",
                        "content": f"The user has invoked the '{skill.name}' skill. Follow these instructions:\n\n{skill.instructions}",
                    }
                )

                # Also add a user message to trigger the model
                self.conversation_history.append(
                    {"role": "user", "content": f"Execute the {skill.name} skill."}
                )
            else:
                # Not a valid skill command, treat as normal user input
                self.conversation_history.append({"role": "user", "content": user_input})
        else:
            # Normal user input (not a skill command)
            self.conversation_history.append({"role": "user", "content": user_input})

        # Stream response from Ollama
        response_text = ""
        self.interrupted = False
        old_handler = None

        def interrupt_handler(signum, frame):
            self.interrupted = True

        with Live(
            Spinner("dots", text="[cyan]baking...[/cyan]"), console=console, transient=True
        ) as live:
            try:
                # Chat with tools - always pass them
                stream = ollama.chat(
                    model=self.model_name,
                    messages=self.conversation_history,
                    tools=TOOLS,
                    stream=True,
                )

                # Install interrupt handler for streaming
                old_handler = signal.signal(signal.SIGINT, interrupt_handler)

                # Collect + render the response (as Markdown) and check for tool calls
                response_text, tool_calls = self._stream_assistant_text(stream, spinner_live=live)

                # Make sure spinner is stopped
                if live.is_started:
                    live.stop()

                # DON'T restore handler yet - keep it active for tool execution

                if self.interrupted:
                    # Restore handler before returning
                    if old_handler:
                        signal.signal(signal.SIGINT, old_handler)
                    # Add partial response to history
                    if response_text:
                        self.conversation_history.append(
                            {"role": "assistant", "content": response_text}
                        )
                    return False

                # If model provided a text response, show it
                if response_text:
                    console.print()  # Newline after response

                # Debug: Check if we got anything at all
                if not response_text and not tool_calls:
                    # Restore handler for retry
                    if old_handler:
                        signal.signal(signal.SIGINT, old_handler)
                    console.print("[dim]Model doesn't support tools, retrying without...[/dim]")

                    # Retry without tools
                    retry_stream = ollama.chat(
                        model=self.model_name, messages=self.conversation_history, stream=True
                    )

                    response_text, _ = self._stream_assistant_text(retry_stream)

                    console.print()  # Newline

                    if not response_text:
                        console.print(
                            "[yellow]⚠️  Model returned empty response even without tools[/yellow]"
                        )
                        return True

                # Add assistant response to history
                assistant_message = {
                    "role": "assistant",
                    "content": response_text,  # Ensure it's the actual string, not wrapped
                }
                if tool_calls:
                    assistant_message["tool_calls"] = tool_calls
                self.conversation_history.append(assistant_message)

                # Handle tool calls, looping so the model can run multiple rounds
                # of tools (e.g. run a command, see its output, then fetch a page
                # and build a table) instead of being cut off after a single round.
                # MAX_TOOL_ROUNDS guards against infinite loops.
                rounds = 0
                while tool_calls and rounds < self.MAX_TOOL_ROUNDS:
                    rounds += 1
                    console.print()  # Newline before tool execution
                    for tool_call in tool_calls:
                        # Check for interrupt before each tool call
                        if self.interrupted:
                            console.print("\n[yellow]⚠️  Tool execution interrupted[/yellow]")
                            # Restore handler and return
                            if old_handler:
                                signal.signal(signal.SIGINT, old_handler)
                            return False

                        try:
                            function = tool_call.get("function", {})
                            name = function.get("name")
                            arguments = function.get("arguments", {})

                            result = self._dispatch_tool_call(name, arguments)

                            # Add tool result to history
                            self.conversation_history.append({"role": "tool", "content": result})

                            # Check for interrupt after the tool ran
                            if self.interrupted:
                                if old_handler:
                                    signal.signal(signal.SIGINT, old_handler)
                                return False

                        except Exception as e:
                            console.print(f"[red]Error executing tool: {e}[/red]")
                            import traceback

                            traceback.print_exc()
                            # Feed the error back so the model can react/correct
                            self.conversation_history.append(
                                {"role": "tool", "content": f"Error executing tool: {e}"}
                            )

                    # Check for interrupt after this round's tools
                    if self.interrupted:
                        console.print("\n[yellow]⚠️  Interrupted before final response[/yellow]")
                        # Restore handler and return
                        if old_handler:
                            signal.signal(signal.SIGINT, old_handler)
                        return False

                    # Get the model's next response given the tool results. It may
                    # contain MORE tool calls (→ another round) or a final text
                    # answer (→ tool_calls is None and the loop ends).
                    next_text = ""
                    tool_calls = None
                    with Live(
                        Spinner("dots", text="[cyan]baking...[/cyan]"),
                        console=console,
                        transient=True,  # Clean up after completion
                    ) as final_live:
                        final_stream = ollama.chat(
                            model=self.model_name,
                            messages=self.conversation_history,
                            tools=TOOLS,  # Include tools so model knows context
                            stream=True,
                        )

                        # Collect + render this response as Markdown; capture any
                        # further tool calls so the while loop can run another round.
                        next_text, tool_calls = self._stream_assistant_text(
                            final_stream, spinner_live=final_live
                        )

                        # Make sure spinner is stopped
                        if final_live.is_started:
                            final_live.stop()

                    if next_text:
                        console.print()  # Newline after any text

                    # Record this response (with any further tool calls) in history
                    assistant_message = {"role": "assistant", "content": next_text}
                    if tool_calls:
                        assistant_message["tool_calls"] = tool_calls
                    self.conversation_history.append(assistant_message)

                    # If interrupted mid-stream, stop here
                    if self.interrupted:
                        if old_handler:
                            signal.signal(signal.SIGINT, old_handler)
                        return False

                    # If tool_calls is set, the while loop runs another round.

                if tool_calls and rounds >= self.MAX_TOOL_ROUNDS:
                    console.print(
                        f"[yellow]⚠️  Stopped after {self.MAX_TOOL_ROUNDS} tool rounds "
                        "(possible loop). The model may not have finished.[/yellow]"
                    )

                # Restore original handler at the very end
                if old_handler:
                    signal.signal(signal.SIGINT, old_handler)

            except KeyboardInterrupt:
                # Restore handler
                if old_handler:
                    signal.signal(signal.SIGINT, old_handler)
                live.stop()
                console.print("\n[yellow]⚠️  Generation interrupted[/yellow]")
                self.interrupted = True
                return False

            except Exception as e:
                # Restore handler
                if old_handler:
                    signal.signal(signal.SIGINT, old_handler)
                live.stop()
                console.print(f"[red]Error: {e}[/red]")
                return False

        return not self.interrupted

    def clear_history(self):
        """Clear conversation history."""
        self.conversation_history = []
